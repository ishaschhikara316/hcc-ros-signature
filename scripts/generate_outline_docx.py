#!/usr/bin/env python3
"""Generate structural overview .docx for the 11-gene oxidative stress-redox HCC paper."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "docs", "structural_overview.docx")
os.makedirs(os.path.join(BASE, "docs"), exist_ok=True)

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Helper functions
def add_heading_1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_heading_2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 76, 153)

def add_heading_3(text):
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.size = Pt(11)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_para(text):
    doc.add_paragraph(text)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("STRUCTURAL OVERVIEW")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "An 11-Gene Oxidative Stress–Redox Prognostic Signature Predicts\n"
    "Overall Survival in Hepatocellular Carcinoma:\n"
    "Multi-Cohort Validation and Biological Characterisation"
)
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Isha Chhikara")
run.font.size = Pt(12)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("March 2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 1. TITLE PAGE AND ABSTRACT
# ═══════════════════════════════════════════════════════════════

add_heading_1("1. Title Page and Abstract")

add_bullet("Proposed title: \"An 11-Gene Oxidative Stress\u2013Redox Prognostic Signature Predicts Overall Survival in Hepatocellular Carcinoma: Multi-Cohort Validation and Biological Characterisation\"")
add_bullet("Structured abstract covering: Background (HCC heterogeneity, oxidative stress and redox biology), Methods (LASSO-Cox on 75 curated genes, 302 TCGA-LIHC patients, 3 validated external cohorts, 532 patients), Results (C-index 0.700, pooled HR 1.70, p=1.3e-6, I\u00b2=0%), and Conclusions (NRF2-driven signature with clinical and therapeutic implications).")
add_bullet("Keywords: hepatocellular carcinoma, reactive oxygen species, oxidative stress, redox, NRF2-KEAP1, prognostic signature, LASSO, overall survival, ferroptosis.")


# ═══════════════════════════════════════════════════════════════
# 2. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

add_heading_1("2. Introduction")

add_heading_2("2.1 Global Burden of Hepatocellular Carcinoma")
add_bullet("Brief description of HCC incidence (~906,000 new cases/year), mortality, and clinical heterogeneity across aetiologies (HBV, HCV, NAFLD, alcohol).")
add_bullet("Poor survival in advanced disease (5-year OS <20% for late-stage) and the inadequacy of existing staging systems (BCLC, TNM) for molecular stratification.")
add_bullet("Need for robust molecular prognostic markers that capture tumour biology beyond histopathological staging.")
add_bullet("Introduction of oxidative stress and ferroptosis as central features of HCC biology and outcome.")

add_heading_2("2.2 Reactive Oxygen Species and Ferroptosis as a Unifying Theme")
add_bullet("Overview of intratumoral ROS generation from mitochondrial dysfunction, NADPH oxidases, and metabolic reprogramming in HCC.")
add_bullet("Description of ferroptosis as an iron-dependent, lipid peroxidation-driven cell death pathway distinct from apoptosis and necroptosis.")
add_bullet("The dual role of ROS: low levels promote tumour proliferation via signalling; excessive levels trigger ferroptosis and cell death.")
add_bullet("The NRF2-KEAP1 axis as the master regulator: NRF2 activation drives antioxidant defence, glutathione synthesis, and iron homeostasis, creating a survival advantage for cancer cells.")
add_bullet("Contrast between physiological ROS homeostasis and pathological ROS exploitation by tumours.")
add_bullet("Rationale for classifying tumour aggressiveness using oxidative stress and redox biology: tumours that upregulate antioxidant defence (high NRF2 activity) are more aggressive because they resist oxidative cell death.")

add_heading_2("2.3 The NRF2-KEAP1 Pathway as a Natural Framework")
add_bullet("Introduction of the NRF2-KEAP1 pathway as a well-characterised regulatory network governing cellular responses to oxidative stress.")
add_bullet("Summary of key downstream effectors: thioredoxin system (TXNRD1), glutathione metabolism (GSR, SLC7A11, G6PD), heme catabolism (HMOX1), autophagy (SQSTM1/p62), and transcriptional regulation (MAFG, BACH1).")
add_bullet("Evidence that KEAP1/NFE2L2 mutations are recurrent in HCC (5\u201310%) and associated with aggressive phenotype and therapy resistance.")
add_bullet("Conceptual link: the NRF2 transcriptional programme represents a biologically coherent gene set that can be exploited for prognostic modelling.")

add_heading_2("2.4 Central Hypothesis and Objectives")
add_bullet("Statement of the hypothesis that oxidative stress and redox pathway genes carry prognostic information in HCC because tumours that master oxidative stress resistance are inherently more aggressive.")
add_bullet("Primary objective: derive and validate a parsimonious gene-based risk model from curated oxidative stress and ferroptosis pathway genes using LASSO-penalised Cox regression.")
add_bullet("Secondary objectives: (i) characterise pathway biology of risk groups via genome-wide GSEA, (ii) profile immune microenvironment differences, (iii) assess drug sensitivity patterns including ferroptosis inducers, (iv) construct a clinical nomogram integrating molecular and clinical variables.")


# ═══════════════════════════════════════════════════════════════
# 3. UNDERSTANDING ROS, FERROPTOSIS, AND NRF2 BIOLOGY
# ═══════════════════════════════════════════════════════════════

add_heading_1("3. Understanding ROS, Ferroptosis, and NRF2 Biology")

add_heading_2("3.1 Cellular and Molecular Basis of Reactive Oxygen Species")
add_bullet("Definition of ROS (superoxide, hydrogen peroxide, hydroxyl radicals) and their generation from mitochondrial electron transport, NADPH oxidases (NOX/NCF2), and metabolic enzymes.")
add_bullet("Description of the antioxidant defence systems: thioredoxin (TXNRD1, TXN), glutaredoxin (GLRX2), glutathione (GSR, GCLC, GCLM), catalase, and superoxide dismutases.")
add_bullet("The NADPH supply chain: G6PD (pentose phosphate pathway) as the critical source of reducing equivalents for all major antioxidant systems.")
add_bullet("Explanation of how these systems are adaptive in normal physiology but become exploited by tumour cells to resist oxidative cell death.")

add_heading_2("3.2 Ferroptosis: Iron-Dependent Cell Death")
add_bullet("Mechanistic overview of ferroptosis: iron-catalysed lipid peroxidation of polyunsaturated fatty acid-containing phospholipids.")
add_bullet("Key regulators: GPX4 (primary ferroptosis suppressor), SLC7A11/xCT (cystine import for glutathione synthesis), ACSL4/LPCAT3 (lipid substrate generation), and NCOA4 (ferritinophagy).")
add_bullet("The role of HMOX1 in ferroptosis: heme degradation releases free iron, creating a pro-ferroptotic environment, yet HMOX1 is paradoxically upregulated by NRF2 as a stress response.")
add_bullet("Ferroptosis sensitivity as a therapeutic vulnerability: erastin (SLC7A11 inhibitor) and RSL3 (GPX4 inhibitor) as prototype ferroptosis inducers.")

add_heading_2("3.3 The NRF2-KEAP1 Signalling Axis")
add_bullet("Detailed description of KEAP1-mediated ubiquitination and proteasomal degradation of NRF2 under basal conditions.")
add_bullet("Mechanisms of NRF2 activation: KEAP1 cysteine modification by electrophiles/ROS, SQSTM1/p62-mediated KEAP1 sequestration, and somatic KEAP1/NFE2L2 mutations in cancer.")
add_bullet("Transcriptional programme of NRF2: antioxidant response element (ARE)-driven expression of TXNRD1, G6PD, SLC7A11, GSR, HMOX1, GCLC, NQO1, and others.")
add_bullet("The MAFG-BACH1 axis: MAFG as obligate heterodimerisation partner for NRF2, BACH1 as transcriptional repressor competing for ARE binding; dynamic balance determines target gene expression.")

add_heading_2("3.4 Oxidative Stress and Ferroptosis Gene Sets and Functional Categories")
add_bullet("Description of the curated 75-gene ROS/ferroptosis gene set assembled from MSigDB Hallmark ROS pathway, KEGG ferroptosis, Reactome detoxification of ROS, and literature-derived NRF2 targets.")
add_bullet("Grouping of genes into functional modules: (i) antioxidant defence (TXNRD1, GSR, GLRX2, G6PD), (ii) NRF2 transcriptional regulation (MAFG, BACH1, SQSTM1), (iii) ferroptosis regulation (SLC7A11, HMOX1), (iv) ROS production and repair (NCF2, primarily immune cell-derived, reflecting tumor microenvironment; MSRA).")
add_bullet("Justification of the relevance of these modules to both normal redox homeostasis and tumour biology in HCC.")


# ═══════════════════════════════════════════════════════════════
# 4. ROS/FERROPTOSIS AND CANCER: CONCEPTUAL FRAMEWORK
# ═══════════════════════════════════════════════════════════════

add_heading_1("4. Oxidative Stress, Redox Biology, and Cancer: Conceptual Framework")

add_heading_2("4.1 Oxidative Stress in Solid Tumours")
add_bullet("Overview of how metabolic reprogramming (Warburg effect), oncogene activation, and mitochondrial dysfunction elevate baseline ROS in tumours.")
add_bullet("Description of consequences: genomic instability, EMT promotion, angiogenesis stimulation, and therapy resistance.")
add_bullet("Note that HCC is particularly ROS-rich due to chronic inflammation (hepatitis), iron overload, cirrhotic microenvironment, and high metabolic activity of hepatocytes.")

add_heading_2("4.2 Ferroptosis Resistance as a Hallmark of Aggressive Tumours")
add_bullet("Presentation of the concept: tumours that upregulate ferroptosis defence (high SLC7A11, high NRF2 activity, high glutathione) survive oxidative challenges and are more aggressive.")
add_bullet("Evidence from HCC: KEAP1 loss-of-function mutations activate NRF2, conferring resistance to sorafenib and other therapies.")
add_bullet("Placement of HCC within ferroptosis-resistant cancers with strong NRF2 and antioxidant signatures correlating with poor outcome.")

add_heading_2("4.3 From Redox Biology to Prognostic Modelling")
add_bullet("Conceptualisation of oxidative stress and redox pathway genes as a biologically curated feature space for survival prediction.")
add_bullet("Description of how tumour co-option of antioxidant and iron-handling pathways creates a transcriptional signature of aggressiveness.")
add_bullet("Positioning of this work as a bridge between redox biology, ferroptosis research, and clinical prognostics in HCC.")
add_bullet("Advantage over transcriptome-wide approaches: every gene in the signature has a clear mechanistic role in oxidative stress and redox biology, enabling biological interpretation.")


# ═══════════════════════════════════════════════════════════════
# 5. STUDY AIMS AND DESIGN OVERVIEW
# ═══════════════════════════════════════════════════════════════

add_heading_1("5. Study Aims and Design Overview")

add_heading_2("5.1 Overall Study Design")
add_bullet("Outline of the strategy: curation of 75 ROS/ferroptosis genes from 4 pathway databases, construction of a LASSO-Cox prognostic model in TCGA-LIHC (n=302), and validation in 3 independent cohorts (532 patients) spanning 3 continents (GSE76427 OS excluded due to insufficient power).")
add_bullet("Emphasis on overall survival as the primary endpoint and risk stratification into high versus low risk using the training-set median cutoff.")
add_bullet("Note on the multi-cohort, multi-platform (RNA-seq and microarray) study design ensuring cross-platform generalisability.")
add_bullet("Secondary endpoints: recurrence-free survival (where available), pathway characterisation, immune profiling, and drug sensitivity.")

add_heading_2("5.2 Specific Research Questions")
add_bullet("Question 1: Do ROS/ferroptosis pathway genes associate with overall survival in HCC beyond chance? (Addressed by univariate screening + permutation testing.)")
add_bullet("Question 2: Can a parsimonious LASSO-derived signature stratify risk independently of clinical stage, age, and sex? (Addressed by multivariate Cox + interaction tests.)")
add_bullet("Question 3: Do metabolic, immune, and drug-response profiles differ systematically between risk groups? (Addressed by GSEA, ssGSEA, and GDSC analysis.)")
add_bullet("Question 4: Does the signature provide clinically actionable risk stratification when combined with standard staging? (Addressed by nomogram + decision curve analysis.)")
add_bullet("Question 5: Is the prognostic signal reproducible across ethnically diverse, multi-platform validation cohorts? (Addressed by meta-analysis with heterogeneity assessment.)")


# ═══════════════════════════════════════════════════════════════
# 6. METHODS
# ═══════════════════════════════════════════════════════════════

add_heading_1("6. Methods")
add_note("[High-level overview in main text; detailed methodology in dedicated Methodology Section Blueprint below.]")

add_heading_2("6.1 Data Sources and Cohorts")
add_bullet("Training cohort: TCGA-LIHC (n=302, 129 OS events), RNA-seq (FPKM), downloaded via GDC API.")
add_bullet("Validation cohort 1: GSE14520 (n=221, 85 events), Chinese HBV-HCC, Affymetrix microarray.")
add_bullet("Validation cohort 2: ICGC LIRI-JP (n=231, 43 events), Japanese HCC, RNA-seq normalised counts.")
add_bullet("Validation cohort 3: GSE54236 (n=80, 80 events), European HCC, Agilent microarray.")
add_bullet("Validation cohort 4: GSE76427 (n=115/108, 23 OS events / 48 RFS events), European HCC, Illumina microarray.")
add_bullet("Supplementary data: Human Protein Atlas (IHC + single-cell RNA-seq), cBioPortal (somatic mutations), GDSC2 (drug sensitivity), MSigDB/KEGG/Reactome (pathways).")

add_heading_2("6.2 Gene Curation and Preprocessing")
add_bullet("Curation of 75 ROS/ferroptosis genes from 4 sources: MSigDB Hallmark Reactive Oxygen Species Pathway, KEGG Ferroptosis, Reactome Detoxification of ROS, and literature-curated NRF2-KEAP1 targets.")
add_bullet("Deduplication and mapping to expression platforms; 68/75 genes expressed in TCGA-LIHC after quality filtering.")
add_bullet("Within-cohort z-score normalisation of gene expression; training-set parameters stored for reproducibility.")

add_heading_2("6.3 Statistical Modelling and Validation")
add_bullet("Univariate Cox screening of all 68 expressed genes with Benjamini-Hochberg FDR correction (threshold: FDR < 0.10).")
add_bullet("LASSO-Cox regression (L1-penalised, \u03bb=0.03 selected via 5-fold cross-validation) on 19 FDR-significant genes.")
add_bullet("Bootstrap stability analysis (1,000 iterations) and permutation test (500 iterations) for model robustness.")
add_bullet("Risk score = \u03a3(coefficient_i \u00d7 z-scored expression_i); median cutoff for high/low risk stratification.")
add_bullet("External validation: risk score applied with fixed coefficients, no re-fitting; C-index, HR, log-rank p reported per cohort.")
add_bullet("Meta-analysis: fixed-effect and random-effects (DerSimonian-Laird) pooling of HRs across validation cohorts; Cochran\u2019s Q and I\u00b2 for heterogeneity.")

add_heading_2("6.4 Biological and Translational Analyses")
add_bullet("Genome-wide preranked GSEA across 4,045 pathways (Hallmark, KEGG, Reactome, GO-BP); 1,000 permutations; FDR < 0.25.")
add_bullet("ssGSEA-based immune cell infiltration profiling (23 cell types) and immune checkpoint expression analysis (9 genes) with BH FDR correction.")
add_bullet("Drug sensitivity correlation analysis using GDSC2 IC50 data for ferroptosis inducers (erastin, RSL3) and standard-of-care agents.")
add_bullet("Restricted Mean Survival Time (RMST) analysis as a PH-assumption-free supplement to Cox regression.")
add_bullet("Nomogram construction integrating risk score with stage, age, sex, and grade; decision curve analysis for clinical utility.")


# ═══════════════════════════════════════════════════════════════
# 7. RESULTS
# ═══════════════════════════════════════════════════════════════

add_heading_1("7. Results")

add_heading_2("7.1 Gene Curation and Survival-Associated Oxidative Stress Genes")
add_bullet("68 of 75 curated oxidative stress and ferroptosis pathway genes were expressed in TCGA-LIHC after quality filtering.")
add_bullet("Univariate Cox regression identified 19 genes significantly associated with OS (FDR < 0.10), including TXNRD1 (HR=1.72, p=1.05e-9), MAFG (HR=1.68, p=1.71e-9), and G6PD (HR=1.59, p=9.06e-9).")
add_bullet("Permutation testing confirmed that the enrichment of survival-associated genes exceeded chance expectation (p < 0.001).")
add_bullet("Notably, the master regulators KEAP1 (p=0.26) and NFE2L2 (p=0.47) were not individually prognostic, indicating that downstream effectors carry the prognostic signal rather than the regulators themselves.")

add_heading_2("7.2 Construction of the 11-Gene Oxidative Stress\u2013Redox Signature")
add_bullet("LASSO-Cox regression selected 11 genes at \u03bb=0.03 (5-fold CV C-index=0.687\u00b10.054): 10 risk genes (TXNRD1, MAFG, G6PD, SQSTM1, SLC7A11, GSR, NCF2, HMOX1, GLRX2, BACH1) and 1 protective gene (MSRA).")
add_bullet("Bootstrap stability: all 11 genes selected in 57\u201399% of 1,000 bootstrap iterations (MSRA 99%, SQSTM1 57%).")
add_bullet("7 of 11 genes (64%) are NRF2-KEAP1 pathway genes (including regulators and interactors), providing strong biological coherence. Note: GLRX2 is not a confirmed NRF2 target, and BACH1 is a competitive repressor rather than a direct target.")
add_bullet("Risk score defined as \u03a3(\u03b2_i \u00d7 z_i); MSRA is the only gene with a negative coefficient (\u03b2=\u22120.226), consistent with its role in repairing oxidative protein damage.")

add_heading_2("7.3 Prognostic Performance in the Training Cohort")
add_bullet("C-index: 0.700 (95% bootstrap CI: 0.650\u20130.744).")
add_bullet("Hazard ratio: 3.40 (95% CI: 2.03\u20135.70, log-rank p=3.9\u00d710\u207b\u2076).")
add_bullet("Time-dependent AUC: 1-year 0.769 (0.695\u20130.837), 3-year 0.744 (0.665\u20130.821), 5-year 0.750 (0.655\u20130.839).")
add_bullet("Integrated Brier Score: 0.216; calibration R\u00b2: 0.91 (1-year), 0.84 (3-year), 0.73 (5-year).")
add_bullet("Multivariate Cox: risk score remained independently significant (HR=4.12, p=3.5\u00d710\u207b\u2079) after adjustment for age, sex, and stage.")
add_bullet("Sensitivity analysis: signature significant across all cutpoint percentiles (25th\u201375th), HR ranging from 2.05 to 3.52 (all p < 0.001).")

add_heading_2("7.4 External Validation Across Independent Cohorts")
add_bullet("GSE14520 (Chinese HBV, n=221): C-index 0.596, HR 1.79 (1.18\u20132.70), p=0.004.")
add_bullet("ICGC LIRI-JP (Japanese, n=231): C-index 0.662, HR 1.78 (1.24\u20132.57), p=0.002.")
add_bullet("GSE54236 (European, n=80): C-index 0.621, HR 1.68 (1.04\u20132.72), p=0.007.")
add_bullet("GSE76427 (European, n=115): OS not significant (23 events, estimated power 3%); RFS HR 1.46, p=0.125 (power 26%).")
add_bullet("Meta-analysis (excluding underpowered GSE76427 OS): fixed-effect HR=1.70 (1.37\u20132.10), p=1.3\u00d710\u207b\u2076; random-effects HR=1.70, p=1.3\u00d710\u207b\u2076; I\u00b2=0% (no heterogeneity); Cochran\u2019s Q=0.50, p=0.92.")
add_bullet("Power analysis: GSE76427 OS exclusion justified by Schoenfeld formula (23 events, 3% power for HR=0.49). All other cohorts had estimated power 47\u201377%.")

add_heading_2("7.5 Multivariate Independence and Subgroup Analyses")
add_bullet("Risk score remained significant after adjustment for stage, age, and sex in TCGA-LIHC (HR=4.12, p<0.001) and ICGC LIRI-JP (HR=1.79, p=0.004).")
add_bullet("GSE14520 multivariate: HR=1.57, p=0.056 (marginally significant; TNM stage dominant in this HBV cohort).")
add_bullet("No significant interactions: risk_score \u00d7 sex (LRT p=0.058), risk_score \u00d7 age (p=0.145), risk_score \u00d7 stage (p=0.876) \u2014 signature generalises across all subgroups.")
add_bullet("Proportional hazards verified via time-split analysis: HR ratio <2.0 for splits at 12, 24, and 36 months; weakens beyond 48 months (addressed by RMST).")
add_bullet("Formal Schoenfeld residual test (cox.zph equivalent): all 4 transforms non-significant \u2014 identity (p=0.148), log (p=0.505), rank (p=0.101), KM/Grambsch-Therneau (p=0.925) \u2014 PH assumption formally confirmed.")

add_heading_2("7.6 Restricted Mean Survival Time Analysis")
add_bullet("RMST analysis performed as a PH-assumption-free supplement:")
add_bullet("1-year: low-risk lives 1.7 months longer (95% CI: 1.0\u20132.4, p<0.0001).")
add_bullet("3-year: low-risk lives 8.1 months longer (95% CI: 5.1\u201311.3, p<0.0001).")
add_bullet("5-year: low-risk lives 14.5 months longer (95% CI: 8.6\u201320.2, p<0.0001).")
add_bullet("Provides clinically interpretable survival benefit that does not depend on the PH assumption.")

add_heading_2("7.7 Comparison with Existing Signatures")
add_bullet("Training cohort (TCGA-LIHC): oxidative stress\u2013redox signature (C=0.700) outperforms Buffa hypoxia 15-gene (C=0.610, p=0.012) and MKI67 4-gene (C=0.637, p=0.001); comparable to Hong 8-gene (C=0.702, p=NS).")
add_bullet("Validation cohort (GSE14520): oxidative stress\u2013redox signature (C=0.596) comparable to Buffa (C=0.622), Hong (C=0.591), MKI67 (C=0.594).")
add_bullet("Note: formal bootstrap C-index comparison confirms risk score significantly outperforms stage alone (\u0394C=0.110, p=0.002).")

add_heading_2("7.8 Nomogram and Clinical Utility")
add_bullet("Nomogram model comparison: Risk score alone (C=0.700) < Risk+Stage (C=0.714) < Risk+Stage+Age+Sex (C=0.722) < Full model with Grade (C=0.752).")
add_bullet("Full nomogram improvement over stage alone: \u0394C=0.159 (stage C=0.593).")
add_bullet("Decision curve analysis: nomogram provides net benefit over \"treat all\" and \"treat none\" strategies across clinically relevant threshold probabilities.")
add_bullet("NRI at 3 years: 0.110; IDI at 5 years: 0.288 (risk score vs stage; note: NRI/IDI reported as exploratory).")
add_bullet("LRT-based model comparison: adding risk score to stage yields LRT=45.8, p=1.3\u00d710\u207b\u00b9\u00b9 (\u0394C=0.121); adding clinical variables to risk score yields LRT=10.4, p=0.015 (\u0394C=0.020).")
add_bullet("AIC comparison: risk score (AIC=1038.9) preferred over stage alone (AIC=1077.7).")

add_heading_2("7.9 NRF2 Pathway Biology and Genetic Validation")
add_bullet("NRF2 pathway activity score strongly correlates with risk score: Spearman r=0.708 (p=2.75\u00d710\u207b\u2074\u2077) \u2014 strongest biological validation.")
add_bullet("7/11 signature genes are NRF2-KEAP1 pathway genes (including regulators and interactors).")
add_bullet("NRF2 pathway mutations (KEAP1 or NFE2L2): 12.6% in high-risk vs 4.0% in low-risk (OR=3.48, Fisher p=0.011).")
add_bullet("Extended NRF2 pathway (KEAP1/NFE2L2/BACH1): 14.6% vs 4.0% (OR=4.12, p=0.002).")
add_bullet("TP53 mutations: 40.4% vs 17.9% (OR=3.11, p<0.0001); CTNNB1: 36.4% vs 15.9% (OR=3.03, p<0.0001) \u2014 both survive FDR correction.")

add_heading_2("7.10 Genome-Wide Pathway Enrichment")
add_bullet("1,384 of 4,045 pathways significantly enriched (FDR < 0.25) between risk groups.")
add_bullet("Top enriched in high-risk: mTORC1 signalling (NES=2.76), MYC targets (NES=2.55), reactive oxygen species pathway (NES=2.43), E2F targets, proteasome, ribosome.")
add_bullet("Top enriched in low-risk: bile acid metabolism, fatty acid oxidation, xenobiotic metabolism, complement \u2014 consistent with preserved hepatocyte differentiation.")
add_bullet("Direct pathway validation: KEAP1-NRF2 pathway (p=1.2\u00d710\u207b\u00b9\u2076), ROS pathway (p=1.3\u00d710\u207b\u00b9\u2079), ferroptosis (p=7.6\u00d710\u207b\u2076).")

add_heading_2("7.11 Ferroptosis Marker Correlations")
add_bullet("13 ferroptosis markers assessed for correlation with risk score:")
add_bullet("SLC7A11 (anti-ferroptosis): r=0.635 (FDR=1.9\u00d710\u207b\u00b3\u2074) \u2014 strongest correlation, confirming cystine-import driven ferroptosis resistance in high-risk tumours.")
add_bullet("TFRC (pro-ferroptosis, iron uptake): r=0.283 (FDR=3.6\u00d710\u207b\u2076); FTL and FTH1 (iron storage): r=0.23 and r=0.22 (FDR<0.001) \u2014 iron accumulation in high-risk.")
add_bullet("7/13 ferroptosis markers significantly correlated (FDR<0.05), 7/13 consistent with expected ferroptosis biology direction.")

add_heading_2("7.12 Immune Microenvironment Features of Risk Groups")
add_bullet("ssGSEA immune profiling: mast cells (r=\u22120.157, p=0.006) and Th17 cells (r=\u22120.125, p=0.030) negatively correlated with risk; other cell types NS after FDR correction.")
add_bullet("Immune checkpoint analysis: only SIGLEC15 significantly downregulated in high-risk after FDR correction (FC=0.57, FDR=1.1\u00d710\u207b\u2075). PD-L1, PD-1, CTLA-4, LAG3, TIM-3, TIGIT, IDO1, B7-H3 all NS after FDR.")
add_bullet("T-cell dysfunction and exclusion scores (TIDE): no significant correlation with risk score (all p>0.3), indicating the signature does not predict immunotherapy response.")
add_bullet("Interpretation: high-risk tumours show a modestly immunosuppressed but not immunotherapy-responsive phenotype; SIGLEC15 as a potential novel target warrants further investigation.")

add_heading_2("7.13 Drug Sensitivity Landscape")
add_bullet("Erastin (ferroptosis inducer, SLC7A11 inhibitor): r=0.635 (p<1\u00d710\u207b\u00b3\u2075) \u2014 high-risk tumours are significantly more sensitive to ferroptosis induction.")
add_bullet("Doxorubicin: r=0.36; 5-FU: r=0.22 \u2014 modest positive correlations.")
add_bullet("RSL3 (GPX4 inhibitor): r=\u22120.11 (p=0.053, NS) \u2014 high-risk tumours may resist GPX4-targeted ferroptosis due to enhanced glutathione reserves.")
add_bullet("Sorafenib (standard-of-care): no significant correlation \u2014 signature captures biology orthogonal to sorafenib targets.")
add_bullet("Clinical implication: high-risk patients may benefit from ferroptosis-inducing strategies targeting the SLC7A11/cystine axis rather than GPX4.")

add_heading_2("7.14 Single-Cell and Protein-Level Validation")
add_bullet("Single-cell RNA-seq (HPA): signature genes originate from multiple cell compartments \u2014 hepatocytes (TXNRD1, SQSTM1, GSR, MSRA, GLRX2), immune cells (G6PD in T cells, NCF2 and HMOX1 in Kupffer cells), and stroma (MAFG, SLC7A11).")
add_bullet("Protein validation (HPA IHC): mRNA-protein concordance confirmed for 4/11 genes (TXNRD1, MAFG, GSR, MSRA); 6/11 genes show weak or absent protein detection in HCC tissue \u2014 a known limitation of IHC sensitivity for moderately expressed proteins.")
add_bullet("Interpretation: the signature captures both tumour-intrinsic (hepatocyte) and microenvironment (immune/stromal) biology, consistent with its prognostic power beyond single-cell-type markers.")

add_heading_2("7.15 Molecular Subtypes and Pan-Cancer Analysis")
add_bullet("Consensus clustering (k=2): Cluster 1 (n=41, 13%) with high risk scores and 71% event rate vs Cluster 2 (n=261, 87%) with 38% event rate.")
add_bullet("Pan-cancer validation: TCGA-KIRC validates (C=0.580, HR=1.73, p=0.001); LUSC, STAD, PAAD do not validate (all p>0.3) \u2014 consistent with liver-specific oxidative stress biology.")
add_bullet("KIRC validation explained by shared VHL/HIF/ROS biology in clear cell renal carcinoma.")


# ═══════════════════════════════════════════════════════════════
# 8. BRIDGING REDOX BIOLOGY AND ONCOLOGY
# ═══════════════════════════════════════════════════════════════

add_heading_1("8. Bridging Redox Biology and Oncology")

add_heading_2("8.1 Biological Rationale for the Signature")
add_bullet("Explanation of how NRF2-driven antioxidant defence genes provide a biologically interpretable and mechanistically grounded feature space.")
add_bullet("Emphasis that every signature gene has a clear functional role in ROS metabolism, ferroptosis regulation, or NRF2 signalling \u2014 unlike transcriptome-wide signatures that may lack mechanistic grounding.")
add_bullet("The signature reflects a convergent transcriptional programme of oxidative stress resistance that is activated in aggressive HCC.")

add_heading_2("8.2 NRF2 Addiction and Ferroptosis Evasion in Cancer")
add_bullet("Discussion of how constitutive NRF2 activation (via KEAP1 mutations, SQSTM1 accumulation, or transcriptional upregulation) creates a \"ROS-resistant\" phenotype.")
add_bullet("Consideration of how genes such as TXNRD1, G6PD, SLC7A11, and GSR collectively build a multi-layered ferroptosis defence.")
add_bullet("MSRA as the lone protective gene: its role in repairing methionine sulfoxide residues on oxidised proteins; downregulation in aggressive tumours suggests loss of protein quality control.")

add_heading_2("8.3 Conceptual Framework for Other Cancer Types")
add_bullet("Proposal to extend the oxidative stress\u2013redox approach to other ROS-driven cancers (KIRC validated; pancreatic, lung squamous as candidates).")
add_bullet("Presentation of the broader concept of \"mining redox biology\" for prognostic and therapeutic biomarkers.")
add_bullet("Suggestion that other stress-response programmes (UPR, autophagy, DNA damage response) could serve as templates for similar curated signature approaches.")


# ═══════════════════════════════════════════════════════════════
# 9. CLINICAL TRANSLATION AND RISK-ADAPTED THERAPY
# ═══════════════════════════════════════════════════════════════

add_heading_1("9. Clinical Translation and Risk-Adapted Therapy")

add_heading_2("9.1 Clinical Risk Stratification and Use Cases")
add_bullet("Description of potential clinical applications: surveillance intensity after resection, transplant prioritisation, adjuvant therapy decisions, and clinical trial stratification.")
add_bullet("Nomogram (C=0.752) refines prognosis beyond BCLC/TNM staging (C=0.593), providing an additional 15.9 concordance points.")
add_bullet("The 3-year RMST difference of 8.1 months provides a patient-interpretable metric for shared decision-making.")

add_heading_2("9.2 Therapeutic Implications by Risk Group")
add_bullet("For high-risk tumours: ferroptosis induction via SLC7A11/cystine axis targeting (erastin sensitivity r=0.635); NRF2 inhibitors to dismantle antioxidant defence; metabolic targeting of the glutathione-NADPH axis.")
add_bullet("For low-risk tumours: standard-of-care with potentially lower-intensity surveillance; preserved hepatocyte differentiation suggests better candidacy for liver-directed therapies.")
add_bullet("SIGLEC15 as a potential immunotherapy target specific to high-risk tumours (significantly downregulated, FDR=1.1\u00d710\u207b\u2075).")

add_heading_2("9.3 Drug Sensitivity and Precision Oncology")
add_bullet("Integration of GDSC-derived drug sensitivities with risk stratification identifies erastin as a lead therapeutic candidate for high-risk patients.")
add_bullet("Sorafenib insensitivity of the risk score suggests the signature captures biology orthogonal to current standard-of-care, identifying patients who need alternative strategies.")
add_bullet("Note on the need for functional validation in HCC cell lines and PDX models before clinical translation.")


# ═══════════════════════════════════════════════════════════════
# 10. LIMITATIONS AND FUTURE DIRECTIONS
# ═══════════════════════════════════════════════════════════════

add_heading_1("10. Limitations and Future Directions")

add_heading_2("10.1 Methodological and Data Limitations")
add_bullet("Retrospective design using archival public cohorts; no prospective validation.")
add_bullet("Platform variability (RNA-seq vs microarray) with cohort-level z-score normalisation; formal batch effect correction not applied.")
add_bullet("Protein-level validation limited: only 4/11 genes show concordant IHC expression (HPA); mRNA-protein discordance for HMOX1, NCF2, G6PD, GLRX2, BACH1.")
add_bullet("GSE76427 OS underpowered (23 events, 3% power); formally excluded from meta-analysis with justification.")
add_bullet("GSE14520 multivariate: risk score marginally significant (p=0.056) when stage included, reflecting the dominance of TNM in this HBV-specific cohort.")
add_bullet("Calibration slopes (0.27\u20130.31) indicate cross-platform coefficient shrinkage; model optimised for discrimination (patient ranking) rather than absolute risk prediction. Recalibration-in-the-large recommended for clinical deployment.")
add_bullet("Sample size justification: EPP=11.7 (\u226510 threshold met), expected Harrell shrinkage=0.915 (>0.9), mitigated by LASSO regularisation and external validation.")
add_bullet("TRIPOD+AI (2024) compliance: 17/22 items fully met, 3 partially met, 2 need text additions (missing data handling statement, AI fairness discussion via interaction tests).")

add_heading_2("10.2 Biological and Clinical Open Questions")
add_bullet("Whether signature genes directly drive aggressiveness or serve as markers of broader NRF2-KEAP1 programme activation.")
add_bullet("Untested interactions with underlying aetiology (HBV vs HCV vs NAFLD) and specific systemic treatments (atezolizumab-bevacizumab, lenvatinib).")
add_bullet("Immune checkpoint findings limited to SIGLEC15; no validated immunotherapy response prediction.")
add_bullet("Molecular subtypes: Cluster 1 (high-risk) is small (n=41, 13%), limiting subgroup power.")
add_bullet("PH assumption weakens beyond 48 months; RMST analysis partially addresses this but long-term prognostic utility uncertain.")

add_heading_2("10.3 Next Steps and Experimental Validation")
add_bullet("Development of a clinical RT-qPCR or NanoString panel for the 11-gene signature for FFPE tissue.")
add_bullet("Prospective-retrospective validation using trial cohorts from sorafenib (SHARP/AP), atezolizumab-bevacizumab (IMbrave150), or lenvatinib (REFLECT) studies.")
add_bullet("Mechanistic experiments: (i) CRISPR knockout of TXNRD1/SLC7A11 in HCC cell lines to test ferroptosis sensitivity, (ii) NRF2 inhibitor (ML385) treatment of high-risk-score cell lines, (iii) erastin/sorafenib combination studies.")
add_bullet("Validation of SIGLEC15 as immunotherapy target in HCC patient-derived organoids.")
add_bullet("Extension to NAFLD-HCC cohorts (growing aetiology globally) and prospective biomarker studies.")


# ═══════════════════════════════════════════════════════════════
# 11. CONCLUSION
# ═══════════════════════════════════════════════════════════════

add_heading_1("11. Conclusion")

add_bullet("An 11-gene oxidative stress\u2013redox signature derived from biologically curated pathway genes robustly stratifies HCC prognosis across 3 validated cohorts (532 patients; pooled HR=1.70, p=1.3\u00d710\u207b\u2076, I\u00b2=0%).")
add_bullet("The signature captures NRF2-KEAP1 pathway activation (r=0.708 with NRF2 activity score) and identifies tumours with enhanced antioxidant defence, ferroptosis resistance, and poor outcome.")
add_bullet("Clinical utility demonstrated through a nomogram (C=0.752) and RMST analysis (3-year survival benefit: 8.1 months), providing actionable risk stratification beyond standard staging.")
add_bullet("Ferroptosis induction via SLC7A11 targeting (erastin sensitivity r=0.635) emerges as a therapeutic strategy for high-risk patients, while SIGLEC15 warrants investigation as an immunotherapy target.")
add_bullet("This work bridges redox biology, ferroptosis research, and clinical prognostics, demonstrating that biologically curated gene sets yield interpretable and reproducible prognostic models.")


# ═══════════════════════════════════════════════════════════════
# 12. FIGURES AND TABLES
# ═══════════════════════════════════════════════════════════════

add_heading_1("12. Figures and Tables")

add_heading_2("12.1 Main Figures (Proposed)")
add_bullet("Figure 1: ", "Study design flowchart. ")
add_bullet("Cohort selection, gene curation, model development, validation, and biological characterisation workflow.")
add_bullet("Figure 2: ", "Signature overview. ")
add_bullet("(A) Heatmap of 11-gene expression sorted by risk score, (B) coefficient bar plot, (C) LASSO cross-validation curve, (D) bootstrap stability.")
add_bullet("Figure 3: ", "Training cohort performance. ")
add_bullet("(A) Kaplan-Meier curves, (B) time-dependent ROC with 95% CI, (C) calibration plot, (D) subgroup forest plot.")
add_bullet("Figure 4: ", "External validation and meta-analysis. ")
add_bullet("(A) Forest plot with pooled HR (meta-analysis), (B-D) KM curves for GSE14520, ICGC, GSE54236, (E) power analysis table.")
add_bullet("Figure 5: ", "RMST analysis. ")
add_bullet("(A) RMST difference bar chart across time horizons, (B) KM curves with RMST area shading.")
add_bullet("Figure 6: ", "Nomogram and clinical utility. ")
add_bullet("(A) Points-based nomogram, (B) calibration curves, (C) decision curve analysis, (D) C-index comparison across models.")
add_bullet("Figure 7: ", "NRF2 pathway and biological validation. ")
add_bullet("(A) Risk score vs NRF2 activity scatter (r=0.708), (B) NRF2 pathway mutation enrichment, (C) ferroptosis marker correlations, (D) gene correlation network.")
add_bullet("Figure 8: ", "Genome-wide GSEA and pathway landscape. ")
add_bullet("(A) Hallmark pathway dotplot, (B) top enriched pathways, (C) KEGG/Reactome enrichment.")
add_bullet("Figure 9: ", "Immune microenvironment and drug sensitivity. ")
add_bullet("(A) Immune checkpoint expression, (B) ssGSEA immune barplot, (C) erastin sensitivity scatter, (D) drug sensitivity heatmap.")

add_heading_2("12.2 Supplementary Figures (Proposed)")
add_bullet("Supp Fig 1: Univariate screening volcano plot and forest plot for all 68 genes.")
add_bullet("Supp Fig 2: Sensitivity analysis \u2014 HR and p-value across cutpoint percentiles (25th-75th).")
add_bullet("Supp Fig 3: Interaction tests forest plot (risk \u00d7 sex, age, stage).")
add_bullet("Supp Fig 4: Leave-one-gene-out analysis bar chart.")
add_bullet("Supp Fig 5: Proportional hazards diagnostics (time-split HR analysis).")
add_bullet("Supp Fig 6: Brier score curves (model vs null).")
add_bullet("Supp Fig 7: Protein validation (HPA IHC concordance).")
add_bullet("Supp Fig 8: Single-cell expression dotplot and cell-type heatmap.")
add_bullet("Supp Fig 9: Molecular subtypes (consensus matrix, cluster KM, TMB).")
add_bullet("Supp Fig 10: Pan-cancer validation (KIRC, LUSC, STAD, PAAD).")
add_bullet("Supp Fig 11: Signature comparison on training and validation cohorts.")

add_heading_2("12.3 Essential Tables (Proposed)")
add_bullet("Table 1: Cohort characteristics (n, events, age, sex, stage, aetiology, platform).")
add_bullet("Table 2: The 11-gene signature: gene symbol, full name, LASSO coefficient, function, NRF2 target status, bootstrap selection frequency.")
add_bullet("Table 3: Validation performance summary (C-index with 95% CI, HR with 95% CI, log-rank p, power estimate per cohort).")
add_bullet("Table 4: Meta-analysis results (fixed/random HR, I\u00b2, Q, tau\u00b2) for main and sensitivity analyses.")
add_bullet("Table 5: RMST differences across time horizons with bootstrap 95% CI.")
add_bullet("Table 6: Multivariate Cox models in training and validation cohorts.")
add_bullet("Table 7: Nomogram model comparison (C-index for each variable combination).")
add_bullet("Table 8: NRF2 pathway mutation analysis (individual genes and pooled tests).")
add_bullet("Table 9: Top 20 enriched Hallmark pathways (NES, FDR).")
add_bullet("Table 10: Ferroptosis marker correlations with risk score (Spearman r, FDR).")
add_bullet("Supp Table 1: Complete univariate Cox results for all 68 genes.")
add_bullet("Supp Table 2: FDR-corrected immune checkpoint results.")
add_bullet("Supp Table 3: FDR-corrected mutation analysis (20 driver genes).")
add_bullet("Supp Table 4: Drug sensitivity correlations.")
add_bullet("Supp Table 5: One-SE rule analysis for LASSO parameter selection.")
add_bullet("Supp Table 6: Sensitivity cutpoint analysis results.")


# ═══════════════════════════════════════════════════════════════
# METHODOLOGY SECTION BLUEPRINT
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
title_m = doc.add_paragraph()
title_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_m.add_run("METHODOLOGY SECTION BLUEPRINT")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 51, 102)
doc.add_paragraph()

add_heading_1("M1. Study Design and Reporting")
add_bullet("Multi-cohort retrospective prognostic study in HCC using public datasets (TCGA, GEO, ICGC).")
add_bullet("Adherence to REMARK (Reporting Recommendations for Tumour Marker Prognostic Studies) guidelines where applicable.")
add_bullet("Software environment: Python 3.x; key libraries: lifelines (Cox/KM), scikit-learn (LASSO), pandas, numpy, scipy, seaborn, matplotlib, gseapy.")
add_bullet("Fully reproducible pipeline: 21 numbered scripts, automated data download, cached intermediates.")

add_heading_1("M2. Data Acquisition and Preprocessing")

add_heading_2("M2.1 Training Cohort: TCGA-LIHC")
add_bullet("RNA-seq expression (FPKM) and clinical data retrieved via GDC API for primary tumour samples.")
add_bullet("Inclusion: primary HCC with available OS data; exclusion: recurrent tumours, missing survival.")
add_bullet("Overall survival defined as time from diagnosis to death (event) or last follow-up (censored).")
add_bullet("Final sample: n=302 patients, 129 OS events, median follow-up ~25 months.")

add_heading_2("M2.2 External Validation Cohorts")
add_bullet("GSE14520 (n=221, 85 events): Chinese HBV-HCC, Affymetrix GPL3921; survival from supplementary annotation; probe-to-gene mapping via platform annotation file, highest mean expression per gene retained.")
add_bullet("ICGC LIRI-JP (n=231, 43 events): Japanese HCC, RNA-seq normalised read counts; survival from donor file (days \u2192 months).")
add_bullet("GSE54236 (n=80, 80 events): European HCC, Agilent microarray; recurrence data used as endpoint.")
add_bullet("GSE76427 (n=115 OS / 108 RFS): European HCC, Illumina BeadChip; limited power for OS (23 events).")
add_bullet("Harmonisation: OS endpoint standardised across cohorts; clinical covariates (age, sex, stage) mapped to common schema where available.")

add_heading_2("M2.3 Gene Expression Preprocessing")
add_bullet("Expression filtering: genes with zero expression in >50% of samples excluded.")
add_bullet("Within-cohort z-score normalisation: z_i = (x_i \u2013 mean_cohort) / sd_cohort for each gene.")
add_bullet("Training-set normalisation parameters (mean, SD) stored in model JSON for reproducibility.")
add_bullet("Handling of missing genes in validation: reduced risk score computed with available genes (\u22653/11 required).")

add_heading_1("M3. Gene Set Curation and Feature Selection")

add_heading_2("M3.1 Curation of Oxidative Stress and Ferroptosis Pathway Genes")
add_bullet("Four source databases: (i) MSigDB Hallmark Reactive Oxygen Species Pathway, (ii) KEGG Ferroptosis (hsa04216), (iii) Reactome Detoxification of Reactive Oxygen Species, (iv) literature-curated NRF2-KEAP1 targets and iron metabolism genes.")
add_bullet("Deduplication across sources yielded 75 unique genes.")
add_bullet("Intersection with TCGA-LIHC expression data: 68/75 genes passed quality filters and entered analysis.")

add_heading_2("M3.2 Univariate Survival Screening")
add_bullet("Univariate Cox proportional hazards models fitted for each of 68 genes using z-scored expression and OS.")
add_bullet("Multiple testing correction: Benjamini-Hochberg FDR with threshold FDR < 0.10.")
add_bullet("19 genes passed FDR threshold and entered LASSO model construction.")
add_bullet("Sensitivity checks: Bonferroni correction and permutation analysis (500 iterations) confirmed enrichment beyond chance.")

add_heading_1("M4. Prognostic Model Development")

add_heading_2("M4.1 LASSO-Cox Model Construction")
add_bullet("L1-penalised Cox regression applied to 19 FDR-significant genes in TCGA-LIHC.")
add_bullet("Grid search over 15 penaliser values (0.001\u20130.5) evaluated by 5-fold cross-validated C-index.")
add_bullet("Penaliser \u03bb=0.03 selected: CV C-index 0.687\u00b10.054, mean 11.8 genes selected per fold.")
add_bullet("One-standard-error rule assessment: selected \u03bb falls within 1SE of optimal, confirming appropriate parsimony.")
add_bullet("Final model: 11 genes with non-zero coefficients; largest coefficient: MSRA (\u03b2=\u22120.226), MAFG (\u03b2=+0.184).")

add_heading_2("M4.2 Risk Score Calculation and Grouping")
add_bullet("Risk score = \u03a3_{i=1}^{11} (\u03b2_i \u00d7 z_i), where \u03b2_i = LASSO coefficient and z_i = z-scored expression.")
add_bullet("Training-set median risk score used as cutoff for high/low risk groups.")
add_bullet("Sensitivity analysis: signature significant at all percentile cutoffs (25th\u201375th).")
add_bullet("Normalisation parameters and coefficients remain fixed during all validation analyses.")

add_heading_1("M5. Model Evaluation and Validation")

add_heading_2("M5.1 Performance Metrics in Training Cohort")
add_bullet("Harrell\u2019s C-index with 95% CI via 1,000 bootstrap resamples.")
add_bullet("Time-dependent AUC (Heagerty incident/dynamic) at 1, 3, and 5 years with 500 bootstrap CIs.")
add_bullet("Integrated Brier Score (IPCW-weighted) over 6\u201360 months.")
add_bullet("Calibration: quintile-based observed-vs-predicted plots; calibration slope and intercept reported.")
add_bullet("Kaplan-Meier curves with log-rank test comparing risk groups.")

add_heading_2("M5.2 External Validation Procedure")
add_bullet("Risk score applied to each validation cohort using fixed LASSO coefficients; gene expression z-scored within each cohort independently.")
add_bullet("Training-set median used as cutoff (or cohort median where risk score distributions differ).")
add_bullet("Same metrics computed: C-index (bootstrap CI), HR (Cox univariate), log-rank p.")
add_bullet("Meta-analysis: fixed-effect and random-effects (DerSimonian-Laird) pooling; Cochran\u2019s Q, I\u00b2, and \u03c4\u00b2 for heterogeneity.")
add_bullet("Power analysis: Schoenfeld formula to estimate statistical power for each cohort given observed events and HR.")

add_heading_2("M5.3 Multivariate and Subgroup Analyses")
add_bullet("Multivariate Cox models including risk score + stage + age + sex in cohorts with available covariates.")
add_bullet("Interaction tests: risk_score \u00d7 sex, risk_score \u00d7 age(\u226560), risk_score \u00d7 stage(III/IV) using likelihood ratio tests.")
add_bullet("Proportional hazards assumption: time-split analysis at 12, 24, 36, 48 months; Schoenfeld residual tests.")
add_bullet("RMST analysis at 1, 2, 3, and 5 years with 1,000 bootstrap CIs as PH-free supplement.")

add_heading_2("M5.4 Benchmarking Against Existing Signatures")
add_bullet("Computation of three comparison signatures on TCGA-LIHC and GSE14520: Hong 8-gene HCC signature, Buffa 15-gene hypoxia signature, MKI67 4-gene proliferation signature.")
add_bullet("Unweighted mean z-score method for comparison signatures (fair comparison with unweighted models).")
add_bullet("Bootstrap C-index difference test (1,000 iterations) for formal comparison with significance threshold p<0.05.")

add_heading_1("M6. Biological and Translational Analyses")

add_heading_2("M6.1 Pathway Enrichment Analysis")
add_bullet("Genome-wide preranked GSEA: all genes ranked by Spearman correlation with risk score; tested against Hallmark, KEGG, Reactome, GO-BP (4,045 gene sets total).")
add_bullet("1,000 permutations; FDR < 0.25 threshold (standard GSEA convention).")
add_bullet("Overrepresentation analysis on the 11 signature genes: GO-BP, KEGG, Reactome with FDR correction.")
add_bullet("NRF2 pathway activity score: mean z-score of available NRF2 target genes (9 signature + additional canonical targets).")

add_heading_2("M6.2 Immune Microenvironment Profiling")
add_bullet("ssGSEA for 23 immune cell types using published gene signatures.")
add_bullet("Spearman correlation of each cell type score with risk score; BH FDR correction across 23 tests.")
add_bullet("Immune checkpoint expression: 9 genes (CD274, PDCD1, CTLA4, LAG3, HAVCR2, TIGIT, SIGLEC15, IDO1, CD276) compared between risk groups via Mann-Whitney U test; BH FDR correction across 9 tests.")
add_bullet("T-cell dysfunction/exclusion scores adapted from TIDE methodology.")

add_heading_2("M6.3 Drug Sensitivity Analysis")
add_bullet("GDSC2 IC50 data for liver and HCC cell lines.")
add_bullet("Spearman correlation of IC50 with risk score-based gene signature expression.")
add_bullet("Focus on ferroptosis inducers (erastin, RSL3) and standard-of-care agents (sorafenib, doxorubicin, 5-FU).")

add_heading_2("M6.4 Mutation Landscape Analysis")
add_bullet("Somatic mutation data from cBioPortal (TCGA-LIHC).")
add_bullet("Fisher\u2019s exact test comparing mutation frequencies between high and low risk groups for 20 HCC driver genes.")
add_bullet("BH FDR and Bonferroni correction across all tested genes.")
add_bullet("NRF2 pathway mutation pooling: KEAP1, NFE2L2, and BACH1 tested individually and as pathway aggregate.")

add_heading_2("M6.5 Gene Contribution Analysis")
add_bullet("Leave-one-gene-out analysis: C-index computed after removing each gene individually from the risk score.")
add_bullet("Bootstrap significance test (500 iterations) for \u0394C-index of each gene removal.")
add_bullet("Functional redundancy analysis: pathway-level gene removal (antioxidant defence, NRF2 signalling, ferroptosis regulation, ROS metabolism).")

add_heading_1("M7. Reproducibility and Data Availability")
add_bullet("Full analysis pipeline implemented in 21 Python scripts (scripts/01\u201321), executable sequentially with automated data download.")
add_bullet("All code, trained model (lasso_model.json), and results available at [GitHub repository URL].")
add_bullet("Raw data from public repositories: TCGA (GDC), GEO (NCBI), ICGC (DCC), HPA, cBioPortal, GDSC.")
add_bullet("To apply the risk score to independent datasets: z-score 11 genes within new cohort, multiply by stored coefficients, sum for risk score.")


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

doc.save(OUT)
print(f"Saved: {OUT}")
